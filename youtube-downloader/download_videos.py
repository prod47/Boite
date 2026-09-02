#!/usr/bin/env python3
"""
Telecharge en MP4 (meilleure qualite disponible) une liste de videos YouTube
listees dans un fichier Excel (.xlsx), Word (.docx), CSV ou texte brut.

Usage:
    python download_videos.py --input liens.xlsx --output "/chemin/vers/disque/dur/Videos"

Voir README.md pour l'installation et le guide complet.
"""

import argparse
import csv
import datetime
import re
import sys
from pathlib import Path

YOUTUBE_URL_RE = re.compile(
    r"https?://(?:www\.|m\.|music\.)?"
    r"(?:youtube\.com/(?:watch\?v=|shorts/|live/|embed/)[\w\-]+(?:[^\s\"'<>]*)?"
    r"|youtu\.be/[\w\-]+(?:[^\s\"'<>]*)?)"
)


def extract_urls_from_text(text: str) -> list[str]:
    return YOUTUBE_URL_RE.findall(text)


def extract_urls(input_path: Path) -> list[str]:
    suffix = input_path.suffix.lower()

    if suffix == ".xlsx":
        import openpyxl

        wb = openpyxl.load_workbook(input_path, data_only=True)
        chunks = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        chunks.append(str(cell.value))
        text = "\n".join(chunks)

    elif suffix == ".docx":
        import docx

        doc = docx.Document(input_path)
        chunks = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    chunks.append(cell.text)
        text = "\n".join(chunks)

    elif suffix == ".csv":
        with open(input_path, newline="", encoding="utf-8-sig") as f:
            text = "\n".join(",".join(row) for row in csv.reader(f))

    else:  # .txt ou autre : on lit tel quel
        text = input_path.read_text(encoding="utf-8", errors="ignore")

    urls = extract_urls_from_text(text)

    # dedoublonnage en gardant l'ordre d'apparition
    seen = set()
    ordered = []
    for u in urls:
        if u not in seen:
            seen.add(u)
            ordered.append(u)
    return ordered


SUPPORTED_EXTENSIONS = {".xlsx", ".docx", ".csv", ".txt"}


def find_input_files(folder: Path) -> list[Path]:
    files = [
        p for p in sorted(folder.iterdir())
        if p.is_file()
        and p.suffix.lower() in SUPPORTED_EXTENSIONS
        and not p.name.startswith("~$")  # fichiers verrous crees par Excel/Word quand ouverts
    ]
    return files


def extract_urls_from_folder(folder: Path) -> list[str]:
    files = find_input_files(folder)
    if not files:
        sys.exit(
            f"Aucun fichier .xlsx/.docx/.csv/.txt trouve dans : {folder}\n"
            "Mets ton tableau de liens dans ce dossier puis relance."
        )

    print("Fichier(s) de liens detecte(s) :")
    for f in files:
        print(f"  - {f.name}")

    seen = set()
    ordered = []
    for f in files:
        for u in extract_urls(f):
            if u not in seen:
                seen.add(u)
                ordered.append(u)
    return ordered


# Privilegie le H.264 (avc1) + AAC : le seul codec garanti lisible tel quel
# dans Adobe Premiere Pro sans transcodage. Les tres hautes resolutions
# (1440p/4K) de YouTube ne sont souvent disponibles qu'en VP9/AV1, que
# Premiere refuse parfois d'importer ("unsupported compression type") -
# on accepte donc de plafonner a 1080p pour garantir un import direct.
FORMAT_PREMIERE_COMPATIBLE = (
    "bv*[vcodec^=avc1]+ba[acodec^=mp4a]/b[vcodec^=avc1]/bv*+ba/b"
)


def build_ydl_opts(output_dir: Path, allow_playlist: bool, cookies_from_browser: str | None):
    opts = {
        "format": FORMAT_PREMIERE_COMPATIBLE,
        "merge_output_format": "mp4",
        "outtmpl": str(output_dir / "%(title)s [%(id)s].%(ext)s"),
        "windowsfilenames": True,
        "noplaylist": not allow_playlist,
        "ignoreerrors": False,
        "quiet": False,
        "no_warnings": False,
    }
    if cookies_from_browser:
        # reutilise la session YouTube du navigateur : evite les blocages
        # "confirme que tu n'es pas un robot" et donne acces aux videos
        # avec restriction d'age / non repertoriees
        opts["cookiesfrombrowser"] = (cookies_from_browser,)
    return opts


def download_all(
    urls: list[str],
    output_dir: Path,
    allow_playlist: bool,
    cookies_from_browser: str | None = None,
) -> list[dict]:
    import yt_dlp

    output_dir.mkdir(parents=True, exist_ok=True)
    ydl_opts = build_ydl_opts(output_dir, allow_playlist, cookies_from_browser)

    results = []
    total = len(urls)

    for i, url in enumerate(urls, start=1):
        print(f"\n=== [{i}/{total}] {url} ===")
        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=True)
            title = info.get("title", "?") if info else "?"
            results.append({"url": url, "status": "OK", "titre": title, "erreur": ""})
            print(f"OK : {title}")
        except Exception as exc:  # on continue sur l'erreur suivante
            results.append({"url": url, "status": "ECHEC", "titre": "", "erreur": str(exc)})
            print(f"ECHEC : {exc}", file=sys.stderr)

    return results


def write_log(results: list[dict], output_dir: Path) -> Path:
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    log_path = output_dir / f"journal_telechargement_{timestamp}.csv"
    with open(log_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["url", "status", "titre", "erreur"])
        writer.writeheader()
        writer.writerows(results)
    return log_path


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    source = parser.add_mutually_exclusive_group(required=True)
    source.add_argument(
        "--input", "-i", type=Path,
        help="Fichier contenant les liens YouTube (.xlsx, .docx, .csv ou .txt)",
    )
    source.add_argument(
        "--input-folder", type=Path,
        help="Dossier a surveiller : tous les fichiers .xlsx/.docx/.csv/.txt qu'il "
             "contient sont lus (nom de fichier libre, peut changer a chaque fois)",
    )
    parser.add_argument(
        "--output", "-o", required=True, type=Path,
        help="Dossier de destination (ex: chemin vers le disque dur)",
    )
    parser.add_argument(
        "--playlist", action="store_true",
        help="Autoriser le telechargement de playlists entieres (par defaut : desactive, "
             "un seul lien = une seule video)",
    )
    parser.add_argument(
        "--cookies-from-browser", metavar="NAVIGATEUR", default=None,
        help="Nom du navigateur (chrome, edge, firefox...) dont reutiliser la session "
             "YouTube connectee. Utile en cas d'erreur 'confirme que tu n'es pas un "
             "robot' ou pour les videos avec restriction d'age.",
    )
    args = parser.parse_args()

    if args.input_folder:
        if not args.input_folder.exists():
            sys.exit(f"Dossier introuvable : {args.input_folder}")
        urls = extract_urls_from_folder(args.input_folder)
    else:
        if not args.input.exists():
            sys.exit(f"Fichier introuvable : {args.input}")
        urls = extract_urls(args.input)

    if not urls:
        sys.exit("Aucun lien YouTube trouve dans le(s) fichier(s) fourni(s).")

    print(f"{len(urls)} lien(s) YouTube trouve(s). Destination : {args.output}")

    results = download_all(urls, args.output, args.playlist, args.cookies_from_browser)
    log_path = write_log(results, args.output)

    ok = sum(1 for r in results if r["status"] == "OK")
    echecs = len(results) - ok
    print(f"\nTermine : {ok} reussie(s), {echecs} echec(s).")
    print(f"Journal detaille : {log_path}")

    if echecs:
        sys.exit(1)


if __name__ == "__main__":
    main()
